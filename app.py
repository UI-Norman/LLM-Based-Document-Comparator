import streamlit as st
import google.generativeai as genai
import PyPDF2
import docx
import io
import os
from dotenv import load_dotenv
import difflib
import re
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from io import BytesIO
import json
from datetime import datetime

# Load environment variables
load_dotenv()

# Configure page
st.set_page_config(
    page_title="AI Document Comparator",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for attractive interface
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(45deg, #FF1744, #E91E63, #9C27B0, #673AB7);
        background-size: 300% 300%;
        animation: headerGlow 4s ease-in-out infinite;
        padding: 4rem 2rem;
        border-radius: 25px;
        margin-bottom: 3rem;
        text-align: center;
        color: white;
        box-shadow: 
            0 0 40px rgba(255, 23, 68, 0.4),
            0 20px 60px rgba(156, 39, 176, 0.3),
            inset 0 1px 0 rgba(255, 255, 255, 0.2);
        border: 4px solid transparent;
        background-clip: padding-box;
        position: relative;
        overflow: hidden;
    }
    
    .main-header::before {
        content: '';
        position: absolute;
        top: -2px;
        left: -2px;
        right: -2px;
        bottom: -2px;
        background: linear-gradient(45deg, #FFD700, #FF69B4, #00BFFF, #32CD32);
        background-size: 400% 400%;
        animation: borderGlow 3s ease-in-out infinite;
        border-radius: 27px;
        z-index: -1;
    }
    
    @keyframes headerGlow {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    
    @keyframes borderGlow {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    
    .main-header h1 {
        font-size: 3.2rem;
        font-weight: 900;
        margin: 0 0 0.8rem 0;
        text-shadow: 
            2px 2px 8px rgba(0,0,0,0.6),
            0 0 20px rgba(255, 255, 255, 0.3),
            0 0 40px rgba(255, 23, 68, 0.5);
        color: #FFFFFF;
        letter-spacing: 2px;
        text-transform: uppercase;
        animation: textPulse 2s ease-in-out infinite;
    }
    
    @keyframes textPulse {
        0%, 100% { transform: scale(1); }
        50% { transform: scale(1.05); }
    }
    
    .main-header p {
        font-size: 1.4rem;
        margin: 0;
        font-weight: 500;
        color: #FFF8E1;
        text-shadow: 1px 1px 4px rgba(0,0,0,0.5);
        opacity: 0.95;
        letter-spacing: 1px;
    }
    
    .warning-box {
        background: #FFF8E1;
        color: #F57F17;
        padding: 2rem;
        border-radius: 12px;
        border: 3px solid #FFC107;
        margin: 1rem 0;
        box-shadow: 0 8px 20px rgba(255, 193, 7, 0.3);
    }
    
    .success-box {
        background: #E8F5E8;
        color: #2E7D32;
        padding: 2rem;
        border-radius: 12px;
        border: 3px solid #4CAF50;
        margin: 1rem 0;
        box-shadow: 0 8px 20px rgba(76, 175, 80, 0.3);
    }
    .stDataFrame {
        max-width: 800px;
        overflow-x: auto;
    }
    .stDataFrame table {
        width: 100%;
        table-layout: fixed;
    }
    .stDataFrame th, .stDataFrame td {
        word-wrap: break-word;
        max-width: 200px;
        overflow: hidden;
        text-overflow: ellipsis;
    }
    
    .download-section {
        background: #F5F5F5;
        padding: 2rem;
        border-radius: 15px;
        border: 2px solid #ddd;
        margin: 2rem 0;
        display: block !important;
        visibility: visible !important;
    }
    
    .stButton > button {
        background: linear-gradient(45deg, #FF1744, #E91E63, #9C27B0, #673AB7);
        color: white !important;
        border-radius: 30px;
        border: 3px solid transparent;
        padding: 1rem 2.5rem;
        font-weight: 800;
        font-size: 1.1rem;
        transition: all 0.3s;
        box-shadow: 0 8px 25px rgba(255, 23, 68, 0.4);
        display: block !important;
        visibility: visible !important;
    }
    
    .stButton > button:hover {
        transform: translateY(-3px);
        box-shadow: 0 12px 35px rgba(255, 23, 68, 0.5);
        background: linear-gradient(45deg, #D81B60, #C2185B, #7B1FA2, #512DA8);
        border-color: #FF1744;
    }
    
    .stDownloadButton > button {
        background: linear-gradient(45deg, #2196F3, #03A9F4, #00BCD4, #009688);
        color: white !important;
        border-radius: 15px;
        border: 2px solid #2196F3;
        padding: 0.8rem 1.5rem;
        font-weight: 700;
        font-size: 1rem;
        transition: all 0.3s;
        box-shadow: 0 4px 15px rgba(33, 150, 243, 0.3);
        margin: 0.5rem;
        display: inline-block !important;
        visibility: visible !important;
    }
    
    .stDownloadButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(33, 150, 243, 0.4);
        background: linear-gradient(45deg, #1976D2, #0288D1, #0097A7, #00796B);
    }
    
    .element-container {
        display: block !important;
        visibility: visible !important;
    }
    
    .stMarkdown {
        display: block !important;
        visibility: visible !important;
    }
    
    .stExpander {
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    
    .metric-container {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px solid #e9ecef;
        margin: 1rem 0;
    }
    
    .stFileUploader {
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    
    .stTextArea {
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    
    .stApp .stWidget {
        background: transparent !important;
        border Ogni: none !important;
        box-shadow: none !important;
    }
</style>
""", unsafe_allow_html=True)

# Initialize Gemini AI
@st.cache_resource
def initialize_ai():
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            st.error("❌ Gemini API key not found! Please add it to your .env file")
            return None
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.5-flash')
        return model
    except Exception as e:
        st.error(f"❌ Error initializing AI: {str(e)}")
        return None

# Extract text from PDF
def extract_pdf_text(file):
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"❌ Error reading PDF: {str(e)}")
        return None

# Extract text from Word document
def extract_docx_text(file):
    try:
        doc = docx.Document(io.BytesIO(file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"❌ Error reading Word document: {str(e)}")
        return None

# Extract text from JSON document
def extract_json_text(file):
    try:
        data = json.load(io.BytesIO(file.read()))
        def flatten_json(obj, parent_key='', sep='.'):
            items = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    new_key = f"{parent_key}{sep}{k}" if parent_key else k
                    items.extend(flatten_json(v, new_key, sep).items())
            elif isinstance(obj, list):
                for i, v in enumerate(obj):
                    new_key = f"{parent_key}{sep}{i}" if parent_key else str(i)
                    items.extend(flatten_json(v, new_key, sep).items())
            else:
                items.append((parent_key, str(obj)))
            return dict(items)
        
        flat_data = flatten_json(data)
        text = "\n".join([f"{key}: {value}" for key, value in flat_data.items()])
        return text.strip()
    except Exception as e:
        st.error(f"❌ Error reading JSON: {str(e)}")
        return None

# Extract text based on file type
def extract_text(file):
    if file.type == "application/pdf":
        return extract_pdf_text(file)
    elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        return extract_docx_text(file)
    elif file.type == "text/plain":
        return str(file.read(), "utf-8")
    elif file.type == "application/json":
        return extract_json_text(file)
    else:
        st.error("❌ Unsupported file format! Please upload PDF, Word, TXT, or JSON files.")
        return None

# Compare documents using AI
def compare_documents_ai(original_text, revised_text, model):
    try:
        prompt = f"""
You are an expert document analyst tasked with creating a comprehensive, detailed, side-by-side comparison of two documents. The goal is to clearly identify and present changes in a professional, structured format with prominent headings and subheadings.

### Instructions:

- Analyze the actual content of both documents (including text, tables, headings, subheadings, lists, code snippets, etc.). Do not assume or invent content; only consider what is explicitly present in the documents.
- Compare all sections, headings, subheadings, and content under them, including tables, numbers, lists, or other structured content.
- Identify exactly what is different, what is added, what is removed, and what is modified.
- Use **prominent, professional, and bold headings/subheadings** (e.g., using Markdown `##` for main headings and `###` for subheadings) to clearly indicate sections and where changes occur.
- For each heading and subheading, present differences in a side-by-side format (preferably tables) with the following details on **separate lines** for clarity:
  - **Actual**: The content from the original document.
  - **Revised**: The content from the revised document.
  - **Change Description**: A clear explanation of what has changed.
  - **Change Type**: Specify whether the change is an Addition, Deletion, Modification, or Structural change in comparison with the actual document.
- Provide a detailed comparative analysis for each heading and subheading, including changes in tables, numbers, lists, or other structured content.
- Include an **Executive Summary** highlighting the overall key changes in comparison with the actual document.
- Provide a **Detailed Comparison** section with examples and direct quotes from the documents.
- Categorize changes into:
  - Minor Edits (e.g., typos, small wording changes) in comparison with the actual document.
  - Substantial Revisions (e.g., rephrased sections, significant content changes) in comparison with the actual document.
  - Critical Updates (e.g., changes affecting meaning, purpose, or legal/financial implications) in comparison with the actual document.
- Include an **Impact Assessment** for important changes, describing how each change affects the meaning, purpose, or interpretation, and indicate which version (original or revised) is better or more acceptable.
- Use clear, bold Markdown headings (e.g., `##`, `###`), bullet points, tables, and side-by-side comparisons for readability.
- Ensure consistency in terminology (use "Actual" instead of "Original" throughout the output to align with the provided example).

### Important:

- Do not make assumptions about missing content.
- Do not summarize content without a side-by-side comparison.
- Compare only sections, tables, lists, and content explicitly present in the documents.
- Ensure that headings like "Actual," "Revised," "Change Description," and "Change Type" are bold and prominent in the output.

### Required Output Format Example:

## Executive Summary

[Summary of key changes in comparison with actual document]

## Section-by-Section Comparison

### Heading: [Heading Name]

| Actual | Revised | Change Description | Change Type |
|------------|-------------|------------------------|------------------------------------------------------|
| [text]     | [text]      | [description of what changed] | [Added/Deleted/Modified/Structural change]|

### Subheading: [Subheading Name]

**Actual**: Content from Original Document
**Revised**: Content from Revised Document
**Change Description**: Clear explanation of what changed
**Change Type**: Addition/Deletion/Modification/Structural change 

[Repeat for all headings, subheadings, tables, and sections.]

## Key Changes Identified

- [List most significant changes in comparison with actual document]

## Impact Assessment

- [Implications of changes in comparison with actual document, including which version is better or acceptable]

## Change Categories

- **Minor Edits**: [...] in comparison with actual document
- **Substantial Revisions**: [...] in comparison with actual document
- **Critical Updates**: [...] in comparison with actual document

## Recommendations

- [Suggestions for improvements or points to review in comparison with actual document]
### Important:

- Do not make assumptions about missing content.
- Do not summarize content without a side-by-side comparison.
- Compare only sections, tables, lists, and content explicitly present in the documents.
- **VERY IMPORTANT: When creating tables, ensure every column boundary is marked by a single pipe symbol (|). Keep the text in table cells, especially 'Change Description' and 'Change Type,' short and concise to prevent the table from breaking.**
---

ACTUAL DOCUMENT:
{original_text}

REVISED DOCUMENT:
{revised_text}
        """

        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"❌ Error in AI comparison: {str(e)}")
        return None

# Parse Markdown tables
def parse_markdown_tables(text):
    tables = []
    lines = text.splitlines()
    current_table = None
    in_table = False
    is_header = True

    for line in lines:
        line = line.strip()
        if not line:
            if in_table and current_table and current_table["rows"]:
                tables.append(current_table)
            in_table = False
            current_table = None
            is_header = True
            continue

        # Handle table lines
        if line.startswith("|"):
            # A table separator line (e.g., |---|---|)
            if re.match(r'^\|(\s*:?-+:?\s*\|)+$', line):
                if in_table and is_header:
                    is_header = False  # Confirmed header is complete
                continue  # Skip separator line

            # Check for placeholder lines (like those in your image: |---|---|---|)
            # This pattern matches one or more groups of (non-whitespace characters then a pipe)
            if all(re.match(r'^(-+|_|)+\s*$', cell.strip()) for cell in line.strip("|").split("|") if cell.strip()):
                continue # Skip placeholder/junk line

            if not in_table:
                in_table = True
                is_header = True
                # Extract headers and remove Markdown asterisks
                headers = [re.sub(r'\*+([^*]+)\*+', r'\1', cell.strip()) for cell in line.strip("|").split("|") if cell.strip()]
                # If the line contains only blank cells, it's not a valid header
                if not any(headers):
                    in_table = False
                    continue
                current_table = {"headers": headers, "rows": []}
            # app1.py (Around line 380 - Revised logic)
            elif not is_header:
                # Extract row data, remove all Markdown formatting (asterisks, underscores)
                cells_raw = line.strip("|").split("|")
                cells = []
                for cell in cells_raw:
                    clean_cell = re.sub(r'(\*\*|__|\*|_)([^*_]*)\1', r'\2', cell.strip()) # Clean up standard bold/italics
                    clean_cell = clean_cell.strip()
                    cells.append(clean_cell)

                # Simple count check:
                if len(cells) == len(current_table["headers"]) + 2: # |A|B|C| -> ['', 'A', 'B', 'C', ''] (5 parts for 3 headers)
                    cells = cells[1:-1] # Remove the empty start/end parts
                elif len(cells) == len(current_table["headers"]) + 1 and not cells[0]: # |A|B|C -> ['', 'A', 'B', 'C'] (4 parts for 3 headers)
                    cells = cells[1:]
                elif len(cells) == len(current_table["headers"]) + 1 and not cells[-1]: # A|B|C| -> ['A', 'B', 'C', ''] (4 parts for 3 headers)
                    cells = cells[:-1]
                elif len(cells) != len(current_table["headers"]):
                    # If it's still the wrong length, it's likely a malformed line we should skip.
                    continue
                
                # Final validity check: must have the correct number of columns AND contain actual text
                if len(cells) == len(current_table["headers"]) and any(cells):
                    # One final clean on each cell: remove all leading/trailing whitespace
                    cells = [c.strip() for c in cells]
                    current_table["rows"].append(cells)
    if in_table and current_table and current_table["rows"]:
        tables.append(current_table)

    return tables

# Process Markdown content
def process_markdown(text, output_format="txt"):
    lines = text.splitlines()
    processed_lines = []
    current_heading_level = 0
    in_table = False
    current_table_lines = []

    for line in lines:
        line = line.strip()
        if not line:
            processed_lines.append({"type": "text", "content": ""})
            continue

        # Handle tables
        if line.startswith("|"):
            in_table = True
            current_table_lines.append(line)
            continue
        elif in_table and not line.startswith("|-"):
            in_table = False
            if current_table_lines:
                processed_lines.append({"type": "table", "content": "\n".join(current_table_lines)})
                current_table_lines = []
            if line.startswith("|-"):
                continue

        # Handle headings
        heading_match = re.match(r'^(#+)\s*(.*)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            # Remove Markdown asterisks from heading content
            content = re.sub(r'\*+([^*]+)\*+', r'\1', heading_match.group(2).strip())
            content = re.sub(r'__(.*?)__', r'\1', content)
            processed_lines.append({"type": "heading", "level": level, "content": content})
            continue

        # Handle bold and italic
        line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        line = re.sub(r'__([^_]+)__', r'\1', line)
        line = re.sub(r'\*([^*]+)\*', r'\1', line)
        line = re.sub(r'_([^_]+)_', r'\1', line)

        # Handle lists
        list_match = re.match(r'^[-*]\s+(.*)$', line)
        if list_match:
            content = re.sub(r'\*+([^*]+)\*+', r'\1', list_match.group(1).strip())
            processed_lines.append({"type": "list", "content": content})
            continue

        # Handle plain text
        processed_lines.append({"type": "text", "content": re.sub(r'\*+([^*]+)\*+', r'\1', line)})

    if in_table and current_table_lines:
        processed_lines.append({"type": "table", "content": "\n".join(current_table_lines)})

    return processed_lines

# Generate comparison analysis table
def generate_comparison_table(original_text, revised_text, ai_analysis, original_filename, revised_filename):
    similarity = calculate_similarity(original_text, revised_text)
    length_change = ((len(revised_text) - len(original_text)) / len(original_text) * 100)
    line_diff = len(revised_text.splitlines()) - len(original_text.splitlines())
    word_diff = len(revised_text.split()) - len(original_text.split())

    original_filename = original_filename[:20]
    revised_filename = revised_filename[:20]

    table_data = [
        {
            "Metric": "Document Names",
            "Original": original_filename,
            "Revised": revised_filename,
            "Description": "Compared document names"
        },
        {
            "Metric": "Length Change",
            "Original": f"{len(original_text):,} chars",
            "Revised": f"{len(revised_text):,} chars",
            "Description": f"{length_change:+.1f}% char change"
        },
        {
            "Metric": "Line Difference",
            "Original": f"{len(original_text.splitlines()):,} lines",
            "Revised": f"{len(revised_text.splitlines()):,} lines",
            "Description": f"{line_diff:+,} lines changed"
        },
        {
            "Metric": "Word Difference",
            "Original": f"{len(original_text.split()):,} words",
            "Revised": f"{len(revised_text.split()):,} words",
            "Description": f"{word_diff:+,} words changed"
        },
    ]
    return table_data

# Generate PDF report
def generate_pdf_report(report_text, original_filename, revised_filename, table_data):
    try:
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        y_position = height - 40
        left_margin = 40
        max_width = width - 80

        def draw_wrapped_text(text, x, y, font, size, max_width, leading=15):
            c.setFont(font, size)
            words = text.split()
            current_line = ""
            y_pos = y
            for word in words:
                if c.stringWidth(current_line + word + " ", font, size) < max_width:
                    current_line += word + " "
                else:
                    c.drawString(x, y_pos, current_line.strip())
                    y_pos -= leading
                    current_line = word + " "
                    if y_pos < 50:
                        c.showPage()
                        y_pos = height - 40
                        c.setFont(font, size)
            if current_line.strip():
                c.drawString(x, y_pos, current_line.strip())
                y_pos -= leading
            return y_pos

        def split_text_to_fit(text, font, size, max_width):
            """Split text into lines that fit within max_width."""
            c.setFont(font, size)
            words = text.split()
            lines = []
            current_line = ""
            for word in words:
                if c.stringWidth(current_line + word + " ", font, size) < max_width:
                    current_line += word + " "
                else:
                    lines.append(current_line.strip())
                    current_line = word + " "
            if current_line.strip():
                lines.append(current_line.strip())
            return lines

        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(left_margin, y_position, "Document Comparison Report")
        y_position -= 30

        # Document info
        c.setFont("Helvetica", 12)
        y_position = draw_wrapped_text(f"Original: {original_filename}", left_margin, y_position, "Helvetica", 12, max_width)
        y_position = draw_wrapped_text(f"Revised: {revised_filename}", left_margin, y_position, "Helvetica", 12, max_width)
        y_position = draw_wrapped_text(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", left_margin, y_position, "Helvetica", 12, max_width)
        y_position -= 20

        # Comparison Table
        c.setFont("Helvetica-Bold", 12)
        y_position = draw_wrapped_text("Comparison Analysis Table", left_margin, y_position, "Helvetica-Bold", 12, max_width)
        y_position -= 20  # Add spacing before table

        col_widths = [100, 120, 120, 115]  # Adjusted to fit within max_width (~455 points)
        table_content = [["Metric", "Original", "Revised", "Description"]]
        for row in table_data:
            table_content.append([
                '\n'.join(split_text_to_fit(row["Metric"], "Helvetica", 7, 100)),
                '\n'.join(split_text_to_fit(row["Original"], "Helvetica", 7, 120)),
                '\n'.join(split_text_to_fit(row["Revised"], "Helvetica", 7, 120)),
                '\n'.join(split_text_to_fit(row["Description"], "Helvetica", 7, 115))
            ])

        table = Table(table_content, colWidths=col_widths, rowHeights=[40] * len(table_content))
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('WORDWRAP', (0, 0), (-1, -1), True),
        ]))
        table.wrapOn(c, max_width, y_position)
        table_height = table._height
        table.drawOn(c, left_margin, y_position - table_height)
        y_position -= table_height + 20

        # Process AI analysis
        processed_content = process_markdown(report_text, output_format="pdf")
        for item in processed_content:
            if item["type"] == "heading":
                level = item["level"]
                font_size = 14 if level == 1 else 12 if level == 2 else 10
                c.setFont("Helvetica-Bold", font_size)
                y_position = draw_wrapped_text(item["content"], left_margin, y_position, "Helvetica-Bold", font_size, max_width)
            elif item["type"] == "list":
                c.setFont("Helvetica", 10)
                # Use plain bullet point without Markdown asterisks
                y_position = draw_wrapped_text(f"• {item['content']}", left_margin + 10, y_position, "Helvetica", 10, max_width - 10)
            elif item["type"] == "text" and item["content"]:
                c.setFont("Helvetica", 10)
                y_position = draw_wrapped_text(item["content"], left_margin, y_position, "Helvetica", 10, max_width)
            # REVISED PDF TABLE PROCESSING (Starting around line 600 in app1.py)
            elif item["type"] == "table":
                tables = parse_markdown_tables(item["content"])
                for table_data in tables:
                    table_content = [table_data["headers"]]
                    # Use specific column widths for content
                    col_widths_table = [100, 120, 120, 115]
                    
                    # 1. Prepare table content and calculate required row heights
                    row_heights = [20] # Height for the header row
                    for row in table_data["rows"]:
                        wrapped_row = []
                        max_lines = 1
                        for i, cell in enumerate(row):
                            # Use the specific column width for wrapping
                            lines = split_text_to_fit(cell, "Helvetica", 7, col_widths_table[i])
                            wrapped_row.append('\n'.join(lines))
                            max_lines = max(max_lines, len(lines))
                        
                        table_content.append(wrapped_row)
                        # Estimate height: 10 points for padding/spacing + (lines * font_size)
                        row_heights.append(15 + max_lines * 11.5) 

                    # 2. Create and style the table
                    # Pass the calculated row_heights to the Table constructor
                    table = Table(table_content, colWidths=col_widths_table, rowHeights=row_heights)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 7),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        # ReportLab's word wrapping is limited; our pre-wrapping in `split_text_to_fit` is key.
                    ]))
                    
                    # 3. Draw the table (using wrapOn to confirm height)
                    table.wrapOn(c, max_width, y_position)
                    table_height = table._height
                    
                    if y_position - table_height < 50:
                        c.showPage()
                        y_position = height - 40
                        
                    table.drawOn(c, left_margin, y_position - table_height)
                    y_position -= table_height + 20
            if y_position < 50:
                c.showPage()
                y_position = height - 40

        c.save()
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ Error generating PDF: {str(e)}")
        return None

# Generate DOCX report
def generate_docx_report(report_text, original_filename, revised_filename, table_data):
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Document Comparison Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Document info
        doc.add_paragraph(f"Original Document: {original_filename}")
        doc.add_paragraph(f"Revised Document: {revised_filename}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Comparison Table
        doc.add_heading('Comparison Analysis Table', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Metric", "Original", "Revised", "Description"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        col_widths = [3.0, 4.0, 4.0, 3.0]
        for col_idx, width in enumerate(col_widths):
            table.columns[col_idx].width = Cm(width)

        for row in table_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row["Metric"][:100]
            row_cells[1].text = row["Original"][:150]
            row_cells[2].text = row["Revised"][:150]
            row_cells[3].text = row["Description"][:150]

        # AI Analysis
        doc.add_heading('AI Analysis Results', level=1)
        processed_content = process_markdown(report_text, output_format="docx")
        for item in processed_content:
            if item["type"] == "heading":
                level = min(item["level"], 3)
                doc.add_heading(item["content"], level=level)
            elif item["type"] == "list":
                # Use plain bullet point without Markdown asterisks
                para = doc.add_paragraph()
                para.add_run(f"• {item['content']}")
            elif item["type"] == "text" and item["content"]:
                doc.add_paragraph(item["content"])
            elif item["type"] == "table":
                tables = parse_markdown_tables(item["content"])
                for table_data in tables:
                    table = doc.add_table(rows=1, cols=len(table_data["headers"]))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(table_data["headers"]):
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].runs[0].bold = True
                    for col_idx, width in enumerate(col_widths):
                        table.columns[col_idx].width = Cm(width)
                    for row in table_data["rows"]:
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(row):
                            row_cells[i].text = cell
                            for paragraph in row_cells[i].paragraphs:
                                paragraph.style.font.size = Cm(0.25)

        # Footer
        footer = doc.add_paragraph('Generated by AI Document Comparator')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ Error generating DOCX: {str(e)}")
        return None

# Generate TXT report
def generate_txt_report(report_text, original_filename, revised_filename, table_data, similarity, length_change, original_text, revised_text):
    try:
        table_text = "COMPARISON ANALYSIS TABLE:\n"
        table_text += f"{'Metric':<15} {'Original':<20} {'Revised':<20} {'Description':<25}\n"
        table_text += "-" * 80 + "\n"
        for row in table_data:
            table_text += f"{row['Metric'][:15]:<15} {row['Original'][:20]:<20} {row['Revised'][:20]:<20} {row['Description'][:25]:<25}\n"

        # Process AI analysis tables
        processed_content = process_markdown(report_text, output_format="txt")
        ai_analysis_text = ""
        for item in processed_content:
            if item["type"] == "heading":
                ai_analysis_text += f"\n{item['content']}\n"
            elif item["type"] == "list":
                # Use plain hyphen for bullet points
                ai_analysis_text += f"- {item['content']}\n"
            elif item["type"] == "text" and item["content"]:
                ai_analysis_text += f"{item['content']}\n"
            elif item["type"] == "table":
                tables = parse_markdown_tables(item["content"])
                for table_data in tables:
                    for table_data in tables:
                        headers = table_data["headers"]
                        ai_analysis_text += "\n"
                        ai_analysis_text += f"{headers[0]:<60} {headers[1]:<60} {headers[2]:<80} {headers[3]:<60}\n"
                        ai_analysis_text += "-" * 260 + "\n"
                        for row in table_data["rows"]:
                            ai_analysis_text += f"{row[0]:<60} {row[1]:<60} {row[2]:<80} {row[3]:<60}\n"
        report_text = f"""Document Comparison Report

Original Document: {original_filename}
Revised Document: {revised_filename}
Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Generated By: AI Document Comparator

COMPARISON METRICS:
- Similarity: {similarity}%
- Length Change: {length_change:+.1f}%
- Character Difference: {len(revised_text) - len(original_text):+,}
- Line Difference: {len(revised_text.splitlines()) - len(original_text.splitlines()):+,}

{table_text}
AI ANALYSIS:
{ai_analysis_text}
"""
        return report_text.encode('utf-8')
    except Exception as e:
        st.error(f"❌ Error generating TXT: {str(e)}")
        return None

# Generate JSON report
def generate_json_report(report_text, original_filename, revised_filename, table_data):
    try:
        tables = parse_markdown_tables(report_text)
        processed_content = process_markdown(report_text, output_format="json")
        analysis_content = []
        for item in processed_content:
            if item["type"] == "table":
                table_content = parse_markdown_tables(item["content"])[0]
                analysis_content.append({
                    "type": "table",
                    "headers": table_content["headers"],
                    "rows": table_content["rows"]
                })
            else:
                analysis_content.append(item)

        report_data = {
            "report_metadata": {
                "original_document": original_filename,
                "revised_document": revised_filename,
                "analysis_date": datetime.now().isoformat(),
                "generated_by": "AI Document Comparator"
            },
            "comparison_table": table_data,
            "ai_analysis": analysis_content,
            "document_stats": {
                "analysis_type": "AI-powered document comparison",
                "comparison_method": "AI analysis"
            }
        }
        buffer = BytesIO()
        buffer.write(json.dumps(report_data, indent=4, ensure_ascii=False).encode('utf-8'))
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ Error generating JSON: {str(e)}")
        return None

# Calculate similarity percentage
def calculate_similarity(text1, text2):
    try:
        matcher = difflib.SequenceMatcher(None, text1, text2)
        return round(matcher.ratio() * 100, 1)
    except:
        return 0.0

# Main application
def main():
    st.markdown("""
    <div class="main-header">
        <h1>🤖 AI Document Comparator</h1>
        <p>Upload any two documents and let AI analyze the differences for you!</p>
    </div>
    """, unsafe_allow_html=True)

    model = initialize_ai()
    if not model:
        st.stop()

    # Sidebar
    with st.sidebar:
        st.markdown("### 🔧 System Status")
        if st.button("🔍 Test API Connection"):
            api_key = os.getenv("GEMINI_API_KEY")
            if api_key:
                st.success("✅ API key found")
                try:
                    genai.configure(api_key=api_key)
                    test_model = genai.GenerativeModel('gemini-2.5-flash')
                    test_response = test_model.generate_content("Test connection.")
                    st.success("✅ API connection successful!")
                except Exception as e:
                    st.error(f"❌ API connection failed: {str(e)}")
            else:
                st.error("❌ No API key found")
        
        st.markdown("### 📋 How to Use")
        st.markdown("""
        1. Upload Original Document - First version/baseline document
        2. Upload Revised Document - Updated/modified version
        3. Click Compare - AI analyzes all differences
        4. Review Results - Get comprehensive comparison report
        5. Download Report - Save results in multiple formats
        """)
        
        st.markdown("### 📄 Supported Formats")
        st.markdown("- PDF (.pdf) - All types of PDF documents")
        st.markdown("- Word (.docx) - Microsoft Word documents")
        st.markdown("- Text (.txt) - Plain text files")
        st.markdown("- JSON (.json) - JSON structured data")
        
        st.markdown("### ⚡ Key Features")
        st.markdown("- AI-Powered Analysis - Advanced comparison using Gemini AI")
        st.markdown("- Comprehensive Reports - Detailed change analysis with comparison table")
        st.markdown("- Multiple Export Formats - PDF, DOCX, TXT, JSON")
        st.markdown("- Visual Diff View - Technical line-by-line comparison")
        st.markdown("- Document Preview - Side-by-side text view")

    # Main content area
    col1, col2 = st.columns(2)

    with col1:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.subheader("📄 Original Document")
        original_file = st.file_uploader(
            "Choose the original/baseline document",
            type=['pdf', 'docx', 'txt', 'json'],
            key="original",
            help="Upload the first version of your document for comparison"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.subheader("📄 Revised Document") 
        revised_file = st.file_uploader(
            "Choose the revised/updated document",
            type=['pdf', 'docx', 'txt', 'json'],
            key="revised",
            help="Upload the modified version of your document"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    if original_file and revised_file:
        with st.spinner("🔍 Extracting text from documents..."):
            original_text = extract_text(original_file)
            revised_text = extract_text(revised_file)

        if original_text and revised_text:
            similarity = calculate_similarity(original_text, revised_text)
            length_change = ((len(revised_text) - len(original_text)) / len(original_text) * 100)
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"""
                <div class="success-box">
                    <h4>📄 Original Document</h4>
                    <strong>File:</strong> {original_file.name}<br>
                    <strong>Size:</strong> {len(original_text):,} characters<br>
                    <strong>Lines:</strong> {len(original_text.splitlines()):,}<br>
                    <strong>Words:</strong> {len(original_text.split()):,}
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="success-box">
                    <h4>📄 Revised Document</h4>
                    <strong>File:</strong> {revised_file.name}<br>
                    <strong>Size:</strong> {len(revised_text):,} characters<br>
                    <strong>Lines:</strong> {len(revised_text.splitlines()):,}<br>
                    <strong>Words:</strong> {len(revised_text.split()):,}
                </div>
                """, unsafe_allow_html=True)

            st.markdown(f"""
            <div class="metric-container">
                <h4>📊 Quick Metrics</h4>
                <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 1rem;">
                    <div><strong>Similarity:</strong> {similarity}%</div>
                    <div><strong>Length Change:</strong> {length_change:+.1f}%</div>
                    <div><strong>Size Difference:</strong> {len(revised_text) - len(original_text):+,} chars</div>
                    <div><strong>Line Difference:</strong> {len(revised_text.splitlines()) - len(original_text.splitlines()):+,}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            # Add a text input for the reference name
            reference_name = st.text_input(
                "📝 Enter Reference Name (e.g., Contract V1 to V2)",
                value="refrence", # Set the default value as requested
                key="reference_name_input",
                help="This name will be included in the report header, outside of the comparison tables."
            )
            st.session_state['reference_name'] = reference_name # Store in session state

            if st.button("🔍 COMPARE DOCUMENTS", type="primary", use_container_width=True):
                with st.spinner("🤖 AI is analyzing the differences... This may take a moment."):
                    table_data = generate_comparison_table(
                        original_text,
                        revised_text,
                        "",
                        original_file.name,
                        revised_file.name
                    )
                    st.session_state['table_data'] = table_data
                    
                    ai_analysis = compare_documents_ai(original_text, revised_text, model)
                    if ai_analysis:
                        st.session_state['analysis'] = ai_analysis
                        st.session_state['analysis_date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

            if 'analysis' in st.session_state and 'table_data' in st.session_state:
                st.markdown('<div class="comparison-result">', unsafe_allow_html=True)
                st.markdown("## 📊 Comparison Analysis Table")
                st.dataframe(st.session_state['table_data'], use_container_width=True)
                
                st.markdown("## 🤖 AI Analysis Results")
                st.markdown('<div class="difference-box">', unsafe_allow_html=True)
                st.markdown(st.session_state['analysis'])
                st.markdown('</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("## 📖 Document Preview")
            preview_col1, preview_col2 = st.columns(2)
            with preview_col1:
                st.markdown("### Original Document")
                st.text_area(
                    "Original Content",
                    original_text,
                    height=400,
                    disabled=False,
                    key="original_preview"
                )
            
            with preview_col2:
                st.markdown("### Revised Document") 
                st.text_area(
                    "Revised Content",
                    revised_text,
                    height=400,
                    disabled=False,
                    key="revised_preview"
                )
            
            if 'analysis' in st.session_state and 'table_data' in st.session_state:
                st.markdown("## 💾 Download Analysis Report")
                table_data = st.session_state['table_data']
                ai_analysis = st.session_state['analysis']
                analysis_date = st.session_state.get('analysis_date', 'N/A')
                
                report_text_formatted = f"""Document Comparison Report

Original Document: {original_file.name}
Revised Document: {revised_file.name}
Analysis Date: {analysis_date}
Generated By: AI Document Comparator

COMPARISON METRICS:
- Similarity: {similarity}%
- Length Change: {length_change:+.1f}%
- Character Difference: {len(revised_text) - len(original_text):+,}
- Line Difference: {len(revised_text.splitlines()) - len(original_text.splitlines()):+,}

AI ANALYSIS:
{ai_analysis}
"""
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    txt_buffer = generate_txt_report(
                        ai_analysis,
                        original_file.name,
                        revised_file.name,
                        table_data,
                        similarity,
                        length_change,
                        original_text,
                        revised_text
                    )
                    if txt_buffer:
                        st.download_button(
                            label="📄 Download TXT",
                            data=txt_buffer,
                            file_name=f"comparison_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                
                with col2:
                    pdf_buffer = generate_pdf_report(report_text_formatted, original_file.name, revised_file.name, table_data)
                    if pdf_buffer:
                        st.download_button(
                            label="📄 Download PDF",
                            data=pdf_buffer,
                            file_name=f"comparison_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                
                with col3:
                    docx_buffer = generate_docx_report(report_text_formatted, original_file.name, revised_file.name, table_data)
                    if docx_buffer:
                        st.download_button(
                            label="📄 Download DOCX",
                            data=docx_buffer,
                            file_name=f"comparison_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                
                with col4:
                    json_buffer = generate_json_report(report_text_formatted, original_file.name, revised_file.name, table_data)
                    if json_buffer:
                        st.download_button(
                            label="📄 Download JSON",
                            data=json_buffer,
                            file_name=f"comparison_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                            mime="application/json",
                            use_container_width=True
                        )

    else:
        st.markdown("""
        <div class="warning-box">
            <h4>⏳ Ready to Compare Documents</h4>
            Please upload both documents (original and revised versions) to begin the AI-powered comparison analysis.
            <br><br>
            <strong>Supported file types:</strong> PDF, DOCX, TXT, JSON<br>
            <strong>Analysis includes:</strong> Content changes, structural differences, impact assessment, comparison table, and detailed recommendations.
        </div>
        """, unsafe_allow_html=True)

    st.markdown("""
    <div style="text-align: center; color: #666; margin-top: 2rem; padding: 1rem;">
        <p><strong>🤖 AI Document Comparator</strong> | Powered by Google Gemini AI | Built with Streamlit</p>
        <p>Compare any documents - contracts, reports, articles, code, essays, manuals, JSON data, and more!</p>
        <p><em>Universal document comparison for all your analysis needs</em></p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()